from __future__ import annotations

import io
from datetime import date

from docx import Document

from .models import ContractContext, MatrixRow
from .scoring import risk_score


def to_docx(rows: list[MatrixRow], context: ContractContext) -> bytes:
    document = Document()
    document.add_heading("MAPA DE GERENCIAMENTO DE RISCOS", level=0)
    document.add_paragraph(context.objeto or "Contratação de solução de TIC")
    document.add_paragraph(f"Data: {date.today():%d/%m/%Y}")

    document.add_heading("Contexto da contratação", level=1)
    document.add_paragraph(f"Tipo: {context.tipo_contratacao}")
    document.add_paragraph(f"Criticidade: {context.criticidade}")
    document.add_paragraph(f"Prazo: {context.prazo}")
    document.add_paragraph(f"Modalidade: {context.modalidade}")
    document.add_paragraph(context.contexto)

    document.add_heading("Riscos selecionados", level=1)
    for row in rows:
        if not row.selecionado:
            continue
        document.add_heading(f"{row.id} - {row.risco}", level=2)
        document.add_paragraph(f"Categoria: {row.categoria}")
        document.add_paragraph(f"Causa: {row.causa}")
        document.add_paragraph(f"Consequências: {row.consequencia}")
        document.add_paragraph(
            f"Probabilidade {row.probabilidade} x Impacto {row.impacto} = "
            f"{risk_score(row.probabilidade, row.impacto)} ({row.nivel})"
        )
        document.add_paragraph(f"Estratégia: {row.estrategia}")
        document.add_heading("Ações preventivas", level=3)
        for action in row.acoes_preventivas:
            document.add_paragraph(_format_action(action), style="List Bullet")
        document.add_heading("Ações de contingência", level=3)
        for action in row.acoes_contingencia:
            document.add_paragraph(_format_action(action), style="List Bullet")
        if row.justificativa:
            document.add_paragraph(f"Justificativa: {row.justificativa}")

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _format_action(action) -> str:
    parts = [action.descricao]
    if action.situacao:
        parts.append(action.situacao)
    if action.responsavel:
        parts.append(action.responsavel)
    return " - ".join(part for part in parts if part)
